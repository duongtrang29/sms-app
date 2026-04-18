from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.styles.style import ParagraphStyle
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "docs" / "bao_cao_final.docx"
SOURCE_MD = ROOT / "docs" / "bc_cap_nhat.md"

RED = RGBColor(192, 0, 0)
BLUE = RGBColor(0, 51, 153)

FIGURE_LINES = [
    "Hình 4.1. ERD toàn bộ hệ thống",
    "Hình 4.2. Use Case Diagram tổng quát",
    "Hình 4.3. Activity Diagram: Đăng ký học phần",
    "Hình 4.4. Activity Diagram: Nhập điểm và duyệt điểm",
    "Hình 4.5. Sequence Diagram: Đăng nhập",
    "Hình 4.6. Sequence Diagram: Đăng ký học phần",
    "Hình 5.1. Trang tổng quan quản trị (Admin Dashboard)",
    "Hình 5.2. Trang đăng ký học phần (Sinh viên)",
    "Hình 5.3. Trang nhập điểm (Giảng viên)",
    "Hình 5.4. Trang duyệt và khóa điểm (Admin)",
    "Hình 5.5. Trang báo cáo và thống kê (Admin)",
]

REFERENCE_LINES = [
    "[1] Sommerville, I. Software Engineering, 10th Edition. Pearson Education. 2016.",
    "[2] Beck, K. và cộng sự. Manifesto for Agile Software Development. agilemanifesto.org. 2001.",
    "[3] SEI Carnegie Mellon. CMMI for Development v2.0. Software Engineering Institute. 2018.",
    "[4] IEEE. IEEE Std 830-1998: Recommended Practice for Software Requirements Specifications. IEEE. 1998.",
    "[5] PostgreSQL Global Development Group. PostgreSQL 16 Documentation. postgresql.org. 2024.",
    "[6] Vercel Inc. Next.js 15 Documentation — App Router. nextjs.org. 2024.",
    "[7] Supabase Inc. Supabase Documentation — Database, Auth, Row Level Security. supabase.com. 2024.",
    "[8] shadcn. shadcn/ui Component Library Documentation. ui.shadcn.com. 2024.",
    "[9] [TODO-TEAM: Thêm tài liệu bài giảng GV theo format: Họ tên GV. Tên bài giảng (Chủ đề 1-9). Tên trường. Năm học.]",
]

SCREENSHOT_BLOCKS = [
    {
        "title": "Hình 5.1. Trang tổng quan quản trị (Admin Dashboard)",
        "guide": 'Chèn ảnh: Đăng nhập Admin → /admin/dashboard\nNội dung cần thấy: thẻ tổng số SV, GV, học phần đang mở, yêu cầu phúc khảo',
        "desc": (
            "Hình 5.1 thể hiện trang tổng quan quản trị hệ thống, hiển thị các chỉ số vận hành "
            "quan trọng như số lượng sinh viên đang học, số giảng viên, số học phần đang mở và "
            "số yêu cầu phúc khảo cần xử lý. Giao diện cung cấp các lối tắt đến các chức năng "
            "quản lý thường dùng, giúp quản trị viên theo dõi trạng thái hệ thống một cách "
            "nhanh chóng và tập trung."
        ),
    },
    {
        "title": "Hình 5.2. Trang đăng ký học phần (Sinh viên)",
        "guide": 'Chèn ảnh: Đăng nhập SV → /student/enrollment\nNội dung cần thấy: danh sách học phần mở, cột giảng viên, lịch học, số chỗ còn',
        "desc": (
            "Hình 5.2 mô tả giao diện đăng ký học phần từ phía sinh viên, hiển thị danh sách "
            "các học phần đang mở cùng thông tin giảng viên, lịch học, thời gian đăng ký và "
            "số chỗ còn lại. Tại màn hình này, sinh viên có thể thực hiện đăng ký hoặc hủy "
            "đăng ký trực tiếp, trong khi các điều kiện nghiệp vụ như tiên quyết, trùng lịch, "
            "sĩ số và giới hạn tín chỉ được hệ thống kiểm tra phía máy chủ."
        ),
    },
    {
        "title": "Hình 5.3. Trang nhập điểm (Giảng viên)",
        "guide": 'Chèn ảnh: Đăng nhập GV → /lecturer/grades/[offering-id]\nNội dung cần thấy: bảng điểm từng SV, cột CC/GK/CK, cột tổng tự tính, nút Lưu nháp/Gửi duyệt',
        "desc": (
            "Hình 5.3 thể hiện giao diện nhập điểm dành cho giảng viên, trong đó bảng điểm "
            "liệt kê từng sinh viên của lớp học phần cùng các cột chuyên cần, giữa kỳ, cuối "
            "kỳ và điểm tổng được hệ thống tự động tính toán. Màn hình này hỗ trợ cả thao tác "
            "lưu nháp và gửi duyệt, giúp giảng viên hoàn thiện bảng điểm trước khi chuyển sang "
            "bước phê duyệt chính thức."
        ),
    },
    {
        "title": "Hình 5.4. Trang duyệt và khóa điểm (Admin)",
        "guide": 'Chèn ảnh: Đăng nhập Admin → /admin/grades\nNội dung cần thấy: danh sách điểm cần duyệt, nút Duyệt/Từ chối/Khóa, trạng thái hiện tại',
        "desc": (
            "Hình 5.4 mô tả giao diện quản trị dùng để duyệt và khóa điểm, nơi quản trị viên "
            "xem danh sách các bảng điểm đang chờ xử lý, kiểm tra trạng thái hiện tại và thực "
            "hiện các thao tác duyệt, từ chối hoặc khóa điểm. Đây là màn hình thể hiện rõ quy "
            "trình kiểm soát dữ liệu điểm nhiều bước, bảo đảm tính chính xác trước khi công bố "
            "kết quả cho sinh viên."
        ),
    },
    {
        "title": "Hình 5.5. Trang báo cáo và thống kê (Admin)",
        "guide": 'Chèn ảnh: Đăng nhập Admin → /admin/reports\nNội dung cần thấy: biểu đồ phân bố SV, tỷ lệ qua môn, bảng cảnh báo học vụ',
        "desc": (
            "Hình 5.5 thể hiện module báo cáo tổng hợp, bao gồm các biểu đồ và bảng thống kê "
            "phục vụ quản trị đào tạo như phân bố sinh viên, tỷ lệ qua môn và cảnh báo học vụ. "
            "Thông qua màn hình này, quản trị viên có thể theo dõi nhanh bức tranh tổng quan về "
            "tình hình học tập và đưa ra quyết định quản lý dựa trên dữ liệu được chuẩn hóa từ hệ thống."
        ),
    },
]

TEST_GROUPS = [
    (
        "NHÓM TC01 — XÁC THỰC VÀ PHÂN QUYỀN",
        [
            ("TC01-01", "Đăng nhập đúng tài khoản Admin", "email/pass hợp lệ của Admin", "Redirect /admin/dashboard", "[TODO-TEAM]", "[TODO-TEAM]"),
            ("TC01-02", "Đăng nhập sai mật khẩu", "email đúng, pass sai", "Thông báo lỗi, không redirect", "[TODO-TEAM]", "[TODO-TEAM]"),
            ("TC01-03", "Đăng nhập tài khoản bị khóa (is_active=false)", "email hợp lệ nhưng bị vô hiệu", '"Tài khoản đã bị vô hiệu hóa"', "[TODO-TEAM]", "[TODO-TEAM]"),
            ("TC01-04", "SV truy cập URL /admin/dashboard trực tiếp", "Đã đăng nhập SV, truy cập /admin", "Redirect về /student/dashboard hoặc lỗi 403", "[TODO-TEAM]", "[TODO-TEAM]"),
            ("TC01-05", "GV truy cập URL /student/enrollment", "Đã đăng nhập GV, truy cập /student/enrollment", "Redirect hoặc lỗi 403", "[TODO-TEAM]", "[TODO-TEAM]"),
        ],
    ),
    (
        "NHÓM TC02 — ĐĂNG KÝ HỌC PHẦN",
        [
            ("TC02-01", "Đăng ký hợp lệ (đủ điều kiện, còn chỗ)", "SV đủ điều kiện, HP còn chỗ, đúng thời gian", '"Đăng ký thành công", current_students +1', "[TODO-TEAM]", "[TODO-TEAM]"),
            ("TC02-02", "Đăng ký học phần đã đăng ký trước đó", "SV đã có enrollment REGISTERED cho HP này", '"Bạn đã đăng ký học phần này"', "[TODO-TEAM]", "[TODO-TEAM]"),
            ("TC02-03", "Đăng ký thiếu môn tiên quyết", "HP có tiên quyết, SV chưa qua môn đó", '"Thiếu môn tiên quyết: [tên môn]"', "[TODO-TEAM]", "[TODO-TEAM]"),
            ("TC02-04", "Đăng ký trùng lịch học", "Lịch HP trùng với HP đã đăng ký", '"Trùng lịch với: [tên học phần]"', "[TODO-TEAM]", "[TODO-TEAM]"),
            ("TC02-05", "Đăng ký ngoài thời gian", "enrollment_open chưa đến hoặc đã qua enrollment_close", '"Ngoài thời gian đăng ký học phần"', "[TODO-TEAM]", "[TODO-TEAM]"),
            ("TC02-06", "Đăng ký học phần đã đầy", "current_students = max_students", '"Học phần đã đủ sinh viên"', "[TODO-TEAM]", "[TODO-TEAM]"),
            ("TC02-07", "Hủy đăng ký hợp lệ", "SV có enrollment REGISTERED, đúng thời gian hủy", "status=CANCELLED, current_students -1", "[TODO-TEAM]", "[TODO-TEAM]"),
        ],
    ),
    (
        "NHÓM TC03 — NHẬP ĐIỂM VÀ DUYỆT",
        [
            ("TC03-01", "Nhập điểm hợp lệ, tính điểm tự động", "CC=8, GK=7, CK=6, trọng số 0.1/0.3/0.6", "total_score=6.5, letter_grade=B+, gpa_point=3.0", "[TODO-TEAM]", "[TODO-TEAM]"),
            ("TC03-02", "Nhập điểm âm hoặc lớn hơn 10", "CC=-1 hoặc CK=11", "Validation lỗi, không lưu", "[TODO-TEAM]", "[TODO-TEAM]"),
            ("TC03-03", "GV gửi duyệt, không tự sửa được", "status=SUBMITTED", "GV không chỉnh sửa được, cần Admin unlock", "[TODO-TEAM]", "[TODO-TEAM]"),
            ("TC03-04", "Admin duyệt điểm", "status SUBMITTED → APPROVED", "SV nhìn thấy điểm trong /student/grades", "[TODO-TEAM]", "[TODO-TEAM]"),
            ("TC03-05", "SV cố sửa điểm qua API trực tiếp", "SV gọi UPDATE grades", "RLS từ chối, lỗi permission", "[TODO-TEAM]", "[TODO-TEAM]"),
        ],
    ),
    (
        "NHÓM TC04 — TRIGGER TÍNH ĐIỂM",
        [
            ("TC04-01", "Thay đổi 1 điểm thành phần, trigger tính lại", "Sửa final_score từ 6 → 8", "total_score, letter_grade, gpa_point cập nhật ngay", "[TODO-TEAM]", "[TODO-TEAM]"),
            ("TC04-02", "Thang điểm chữ đúng", "total=3.9 | total=4.5 | total=5.5 | total=7.5 | total=8.5 | total=9.1", "F | D | C | B | A | A+", "[TODO-TEAM]", "[TODO-TEAM]"),
            ("TC04-03", "Lịch sử thay đổi điểm được ghi", "Sửa điểm bất kỳ", "Xuất hiện bản ghi trong grade_change_logs", "[TODO-TEAM]", "[TODO-TEAM]"),
        ],
    ),
    (
        "NHÓM TC05 — PHÚC KHẢO",
        [
            ("TC05-01", "Gửi phúc khảo hợp lệ", "Điểm đang APPROVED, SV chưa có request PENDING", "Tạo regrade_request với status=PENDING", "[TODO-TEAM]", "[TODO-TEAM]"),
            ("TC05-02", "Gửi phúc khảo trùng khi đang PENDING", "Đã có request PENDING cho cùng grade", '"Đã có yêu cầu phúc khảo đang xử lý"', "[TODO-TEAM]", "[TODO-TEAM]"),
            ("TC05-03", "Admin xử lý phúc khảo", "status PENDING → RESOLVED với nội dung giải quyết", "SV nhìn thấy kết quả trong /student/regrade", "[TODO-TEAM]", "[TODO-TEAM]"),
        ],
    ),
]


def configure_base_style(style: ParagraphStyle, size: int = 13, bold: bool = False, italic: bool = False, align: int = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    fmt = style.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.3
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    fmt.alignment = align


def configure_styles(doc: Document) -> None:
    configure_base_style(doc.styles["Normal"], size=13)
    configure_base_style(doc.styles["Heading 1"], size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    configure_base_style(doc.styles["Heading 2"], size=13, bold=True)
    configure_base_style(doc.styles["Heading 3"], size=13, bold=True, italic=True)
    configure_base_style(doc.styles["Caption"], size=13)


def style_paragraph(paragraph: Paragraph, alignment: int | None = None, style_name: str | None = None) -> None:
    if style_name:
        paragraph.style = style_name
    if alignment is not None:
        paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.3
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)


def style_run(run, size_pt: int = 13, bold: bool = False, italic: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")


def add_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def clear_paragraph(paragraph: Paragraph) -> None:
    paragraph._p.clear_content()


def add_toc_field(paragraph: Paragraph) -> None:
    clear_paragraph(paragraph)
    style_paragraph(paragraph, WD_ALIGN_PARAGRAPH.LEFT)
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Cập nhật mục lục trong Word"
    fld_separate.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.extend([fld_begin, instr, fld_separate, fld_end])
    style_run(run, size_pt=13)


def set_cell_border(cell, edge: str, val: str, color: str = "808080", size: str = "10") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    element = tc_borders.find(qn(f"w:{edge}"))
    if element is None:
        element = OxmlElement(f"w:{edge}")
        tc_borders.append(element)
    element.set(qn("w:val"), val)
    element.set(qn("w:sz"), size)
    element.set(qn("w:space"), "0")
    element.set(qn("w:color"), color)


def build_placeholder_table(doc: DocumentObject) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    row = table.rows[0]
    row.height = Cm(8)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = row.cells[0]
    cell.width = Cm(14)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for edge in ("top", "left", "bottom", "right"):
        set_cell_border(cell, edge, "dashSmallGap")
    p = cell.paragraphs[0]
    style_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("[TODO-SCREENSHOT]")
    style_run(run, size_pt=13, bold=True)


def find_line(lines: list[str], pattern: str) -> str:
    regex = re.compile(pattern)
    for line in lines:
        stripped = line.strip()
        if regex.match(stripped):
            return stripped
    raise ValueError(f"Cannot find line for pattern: {pattern}")


def extract_block(lines: list[str], heading: str) -> list[str]:
    start = None
    for i, line in enumerate(lines):
        if line.strip() == heading:
            start = i + 1
            break
    if start is None:
        raise ValueError(f"Heading not found: {heading}")
    end = len(lines)
    current_level = heading.split()[0]
    stop_prefixes = []
    if current_level == "###":
        stop_prefixes = ["### ", "## "]
    elif current_level == "####":
        stop_prefixes = ["#### ", "### ", "## "]
    elif current_level == "##":
        stop_prefixes = ["## "]
    for j in range(start, len(lines)):
        if any(lines[j].startswith(prefix) for prefix in stop_prefixes):
            end = j
            break
    block = lines[start:end]
    while block and not block[0].strip():
        block.pop(0)
    while block and not block[-1].strip():
        block.pop()
    return block


def add_body_lines(doc: DocumentObject, lines: list[str]) -> None:
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("#### "):
            text = line.replace("#### ", "", 1).strip()
            p = doc.add_paragraph()
            style_paragraph(p, WD_ALIGN_PARAGRAPH.LEFT, "Heading 3")
            p.add_run(text)
            continue
        if line.startswith("##### "):
            text = line.replace("##### ", "", 1).strip()
            p = doc.add_paragraph()
            style_paragraph(p, WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(text)
            style_run(run, size_pt=13, bold=True, italic=True)
            continue
        p = doc.add_paragraph()
        align = WD_ALIGN_PARAGRAPH.LEFT if line.startswith("- ") or re.match(r"^\d+\.\s", line) else WD_ALIGN_PARAGRAPH.JUSTIFY
        style_paragraph(p, align)
        run = p.add_run(line)
        style_run(run, size_pt=13)


def add_heading(doc: DocumentObject, text: str, level: int) -> None:
    p = doc.add_paragraph()
    if level == 1:
        style_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER, "Heading 1")
    elif level == 2:
        style_paragraph(p, WD_ALIGN_PARAGRAPH.LEFT, "Heading 2")
    else:
        style_paragraph(p, WD_ALIGN_PARAGRAPH.LEFT, "Heading 3")
    p.add_run(text)


def update_figure_list(doc: DocumentObject) -> None:
    paragraphs = doc.paragraphs
    fig_heading_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "DANH MỤC HÌNH VẼ")
    toc_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "MỤC LỤC")
    first_para = paragraphs[fig_heading_idx + 1]
    clear_paragraph(first_para)
    style_paragraph(first_para, WD_ALIGN_PARAGRAPH.LEFT)
    run = first_para.add_run(FIGURE_LINES[0])
    style_run(run, size_pt=13)
    cursor = first_para
    for line in FIGURE_LINES[1:]:
        cursor = add_paragraph_after(cursor)
        style_paragraph(cursor, WD_ALIGN_PARAGRAPH.LEFT)
        run = cursor.add_run(line)
        style_run(run, size_pt=13)
    # Remove stale placeholder paragraphs until TOC heading.
    updated_paragraphs = doc.paragraphs
    new_toc_idx = next(i for i, p in enumerate(updated_paragraphs) if p.text.strip() == "MỤC LỤC")
    for idx in range(fig_heading_idx + 1 + len(FIGURE_LINES), new_toc_idx):
        clear_paragraph(updated_paragraphs[idx])


def prepare_toc_paragraph(doc: DocumentObject) -> None:
    paragraphs = doc.paragraphs
    toc_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "MỤC LỤC")
    intro_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "LỜI MỞ ĐẦU")
    toc_para = paragraphs[toc_idx + 1] if toc_idx + 1 < intro_idx else add_paragraph_after(paragraphs[toc_idx])
    add_toc_field(toc_para)


def add_screenshot_block(doc: DocumentObject, item: dict[str, str]) -> None:
    build_placeholder_table(doc)
    caption = doc.add_paragraph()
    style_paragraph(caption, WD_ALIGN_PARAGRAPH.CENTER)
    run = caption.add_run(item["title"])
    style_run(run, size_pt=13, bold=True)

    guide = doc.add_paragraph()
    style_paragraph(guide, WD_ALIGN_PARAGRAPH.CENTER)
    parts = item["guide"].split("\n")
    for i, part in enumerate(parts):
        if i:
            guide.add_run().add_break(WD_BREAK.LINE)
        run = guide.add_run(part)
        style_run(run, size_pt=11, italic=True)

    desc = doc.add_paragraph()
    style_paragraph(desc, WD_ALIGN_PARAGRAPH.JUSTIFY)
    run = desc.add_run(item["desc"])
    style_run(run, size_pt=13)


def add_test_case_table(doc: DocumentObject) -> None:
    add_heading(doc, "Bảng 5.1. Kết quả kiểm thử chức năng", 3)
    total_rows = 1 + sum(1 + len(rows) for _, rows in TEST_GROUPS)
    table = doc.add_table(rows=total_rows, cols=6)
    table.style = "Table Grid"
    table.autofit = False
    widths = [1.8, 3.3, 3.0, 4.5, 2.0, 1.4]
    for idx, width in enumerate(widths):
        for cell in table.columns[idx].cells:
            cell.width = Cm(width)

    headers = ["TC-ID", "Mô tả", "Đầu vào", "Kết quả mong đợi", "Kết quả thực tế", "Đạt/Không"]
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = header
        p = cell.paragraphs[0]
        style_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER)
        for run in p.runs:
            style_run(run, size_pt=11, bold=True)

    row_idx = 1
    for group_title, rows in TEST_GROUPS:
        group_cell = table.cell(row_idx, 0)
        merged = group_cell
        for col in range(1, 6):
            merged = merged.merge(table.cell(row_idx, col))
        merged.text = group_title
        p = merged.paragraphs[0]
        style_paragraph(p, WD_ALIGN_PARAGRAPH.LEFT)
        for run in p.runs:
            style_run(run, size_pt=11, bold=True)
        row_idx += 1
        for row in rows:
            for col_idx, value in enumerate(row):
                cell = table.cell(row_idx, col_idx)
                cell.text = value
                p = cell.paragraphs[0]
                align = WD_ALIGN_PARAGRAPH.CENTER if col_idx in (0, 5) else WD_ALIGN_PARAGRAPH.LEFT
                style_paragraph(p, align)
                for run in p.runs:
                    style_run(run, size_pt=11)
            row_idx += 1

    note = doc.add_paragraph()
    style_paragraph(note, WD_ALIGN_PARAGRAPH.LEFT)
    run = note.add_run("[TODO-TEAM: Chạy từng test case thực tế trên hệ thống đang chạy, điền kết quả thực tế và đánh dấu Đạt/Không. KHÔNG điền theo phỏng đoán.]")
    style_run(run, size_pt=13, color=RED)


def apply_heading_styles_to_existing_paragraphs(doc: DocumentObject) -> None:
    heading1_exact = {
        "TRANG PHÂN CÔNG NHIỆM VỤ",
        "DANH MỤC TỪ VIẾT TẮT",
        "DANH MỤC HÌNH VẼ",
        "MỤC LỤC",
        "LỜI MỞ ĐẦU",
        "KẾT LUẬN",
        "TÀI LIỆU THAM KHẢO",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text in heading1_exact or re.match(r"^CHƯƠNG\s+\d+\.", text):
            style_paragraph(paragraph, WD_ALIGN_PARAGRAPH.CENTER, "Heading 1")
            continue
        if re.match(r"^\d+\.\d+\.\d+\.", text):
            style_paragraph(paragraph, WD_ALIGN_PARAGRAPH.LEFT, "Heading 3")
            continue
        if re.match(r"^\d+\.\d+\.", text):
            style_paragraph(paragraph, WD_ALIGN_PARAGRAPH.LEFT, "Heading 2")


def add_reference_lines(doc: DocumentObject) -> None:
    for line in REFERENCE_LINES:
        p = doc.add_paragraph()
        style_paragraph(p, WD_ALIGN_PARAGRAPH.LEFT)
        run = p.add_run(line)
        style_run(run, size_pt=13)


def build_report() -> None:
    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    doc = Document(DOC_PATH)

    if any("CHƯƠNG 5" in p.text for p in doc.paragraphs):
        raise RuntimeError("File đã có Chương 5, dừng để tránh chèn trùng.")

    configure_styles(doc)
    apply_heading_styles_to_existing_paragraphs(doc)

    doc.add_page_break()
    add_heading(doc, "CHƯƠNG 5. CÀI ĐẶT, DEMO, KIỂM THỬ, BẢO TRÌ VÀ ĐÀO TẠO", 1)

    sec_51 = find_line(lines, r"^###\s+5\.1\.")
    add_heading(doc, sec_51.replace("### ", "", 1), 2)
    add_body_lines(doc, extract_block(lines, sec_51))

    add_heading(doc, "5.2. Mô tả một số màn hình demo tiêu biểu", 2)
    for block in SCREENSHOT_BLOCKS:
        add_screenshot_block(doc, block)

    add_heading(doc, "5.3. Kiểm thử phần mềm", 2)
    sec_531 = find_line(lines, r"^####\s+5\.3\.1\.")
    sec_532 = find_line(lines, r"^####\s+5\.3\.2\.")
    add_heading(doc, sec_531.replace("#### ", "", 1), 3)
    add_body_lines(doc, extract_block(lines, sec_531))
    add_heading(doc, sec_532.replace("#### ", "", 1), 3)
    add_body_lines(doc, extract_block(lines, sec_532))
    add_test_case_table(doc)

    for pattern in [r"^###\s+5\.4\.", r"^###\s+5\.5\.", r"^###\s+5\.6\.", r"^###\s+5\.7\.", r"^###\s+5\.8\."]:
        heading = find_line(lines, pattern)
        add_heading(doc, heading.replace("### ", "", 1), 2)
        add_body_lines(doc, extract_block(lines, heading))

    doc.add_page_break()
    add_heading(doc, "KẾT LUẬN", 1)
    conclusion_heading = find_line(lines, r"^##\s+KẾT LUẬN$")
    add_body_lines(doc, extract_block(lines, conclusion_heading))

    doc.add_page_break()
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    add_reference_lines(doc)

    update_figure_list(doc)
    prepare_toc_paragraph(doc)
    apply_heading_styles_to_existing_paragraphs(doc)
    doc.save(DOC_PATH)


if __name__ == "__main__":
    build_report()

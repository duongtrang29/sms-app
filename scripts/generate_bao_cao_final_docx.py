from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ABBREVIATIONS: list[tuple[str, str]] = [
    ("ADMIN", "Quản trị viên hệ thống"),
    ("API", "Application Programming Interface"),
    ("BaaS", "Backend-as-a-Service"),
    ("CNPM", "Công nghệ phần mềm"),
    ("CMM", "Capability Maturity Model"),
    ("CMMI", "Capability Maturity Model Integration"),
    ("DRAFT", "Trạng thái nháp - điểm chưa gửi duyệt"),
    ("SUBMITTED", "Trạng thái đã gửi duyệt"),
    ("APPROVED", "Trạng thái đã duyệt"),
    ("LOCKED", "Trạng thái đã khóa điểm"),
    ("ERD", "Entity-Relationship Diagram"),
    ("FK", "Foreign Key - khóa ngoại"),
    ("GPA", "Grade Point Average - điểm trung bình tích lũy"),
    ("GV", "Giảng viên"),
    ("JWT", "JSON Web Token"),
    ("LECTURER", "Giảng viên - vai trò trong hệ thống"),
    ("PK", "Primary Key - khóa chính"),
    ("RLS", "Row Level Security - bảo mật theo hàng"),
    ("SDLC", "Software Development Life Cycle"),
    ("SMS", "Student Management System"),
    ("SRS", "Software Requirements Specification"),
    ("SV", "Sinh viên"),
    ("STUDENT", "Sinh viên - vai trò trong hệ thống"),
    ("UC", "Use Case - ca sử dụng"),
    ("UI", "User Interface"),
    ("UML", "Unified Modeling Language"),
    ("UX", "User Experience"),
]


def parse_args() -> argparse.Namespace:
    root_dir = Path(__file__).resolve().parents[1]
    parser = argparse.ArgumentParser(
        description="Tạo docs/bao_cao_final.docx (khung 1A) bằng python-docx."
    )
    parser.add_argument(
        "--source-md",
        type=Path,
        default=root_dir / "docs" / "bc_cap_nhat.md",
        help="Đường dẫn file markdown nguồn chứa phần LỜI MỞ ĐẦU.",
    )
    parser.add_argument(
        "--output-docx",
        type=Path,
        default=root_dir / "docs" / "bao_cao_final.docx",
        help="Đường dẫn file docx đầu ra.",
    )
    return parser.parse_args()


def configure_section_page(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.left_margin = Cm(3.0)


def configure_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")

    fmt = style.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.3
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)


def style_paragraph(paragraph, alignment: WD_ALIGN_PARAGRAPH | None = None) -> None:
    if alignment is not None:
        paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.3
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)


def style_run(run, size_pt: int = 13, bold: bool = False, italic: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")


def add_chapter_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    run = paragraph.add_run(text.upper())
    style_run(run, size_pt=14, bold=True)


def add_heading_xx(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, WD_ALIGN_PARAGRAPH.LEFT)
    run = paragraph.add_run(text)
    style_run(run, size_pt=13, bold=True)


def add_heading_xxx(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, WD_ALIGN_PARAGRAPH.LEFT)
    run = paragraph.add_run(text)
    style_run(run, size_pt=13, bold=True, italic=True)


def style_table_text(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                style_paragraph(paragraph, WD_ALIGN_PARAGRAPH.LEFT)
                for run in paragraph.runs:
                    style_run(
                        run,
                        size_pt=13,
                        bold=bool(run.bold),
                        italic=bool(run.italic),
                    )


def extract_loi_mo_dau(source_md: Path) -> str:
    if not source_md.exists():
        raise FileNotFoundError(f"Không tìm thấy file markdown nguồn: {source_md}")

    content = source_md.read_text(encoding="utf-8")
    lines = content.splitlines()

    start_idx = None
    for idx, line in enumerate(lines):
        if line.strip() == "## LỜI MỞ ĐẦU":
            start_idx = idx + 1
            break
    if start_idx is None:
        raise ValueError("Không tìm thấy heading '## LỜI MỞ ĐẦU' trong file markdown.")

    end_idx = len(lines)
    for idx in range(start_idx, len(lines)):
        if lines[idx].startswith("## "):
            end_idx = idx
            break

    block = lines[start_idx:end_idx]
    while block and not block[0].strip():
        block.pop(0)
    while block and not block[-1].strip():
        block.pop()

    if not block:
        raise ValueError("Phần LỜI MỞ ĐẦU rỗng.")

    return "\n".join(block)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_separate.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.extend([fld_begin, instr, fld_separate, fld_end])


def set_section_page_number_start(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    for node in sect_pr.xpath("./w:pgNumType"):
        sect_pr.remove(node)
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:start"), str(start))
    sect_pr.append(pg_num_type)


def add_cover_page(doc: Document) -> None:
    cover_lines = [
        "Trường: [TODO-TEAM: Tên trường]",
        "Khoa: [TODO-TEAM: Tên khoa]",
        "",
        "Môn học: CÔNG NGHỆ PHẦN MỀM",
        "Đề tài: HỆ THỐNG QUẢN LÝ SINH VIÊN",
        "",
        "Nhóm: [TODO-TEAM: Số nhóm]",
        "Giảng viên hướng dẫn: [TODO-TEAM: Họ tên GV]",
        "Năm học: [TODO-TEAM: Năm học]",
    ]

    for line in cover_lines:
        paragraph = doc.add_paragraph()
        style_paragraph(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
        if line:
            run = paragraph.add_run(line)
            if line.startswith("Môn học:") or line.startswith("Đề tài:"):
                style_run(run, size_pt=13, bold=True)
            else:
                style_run(run, size_pt=13, bold=False)


def add_task_assignment_page(doc: Document) -> None:
    add_chapter_title(doc, "TRANG PHÂN CÔNG NHIỆM VỤ")

    headers = ["STT", "Họ và tên", "Mã SV", "Nhiệm vụ đảm nhận", "Mức hoàn thành", "Ký tên"]
    table = doc.add_table(rows=6, cols=6)
    table.style = "Table Grid"
    table.autofit = False

    widths_cm = [1.0, 3.3, 2.2, 4.3, 2.5, 2.7]
    for idx, width in enumerate(widths_cm):
        for cell in table.columns[idx].cells:
            cell.width = Cm(width)

    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = header
        paragraph = cell.paragraphs[0]
        style_paragraph(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
        for run in paragraph.runs:
            style_run(run, size_pt=13, bold=True)

    for row_idx in range(1, 6):
        row = table.rows[row_idx]
        row.cells[0].text = str(row_idx)
        for col_idx in range(1, 6):
            row.cells[col_idx].text = "[TODO-TEAM: điền thông tin]"

    style_table_text(table)

    for row_idx in range(1, 6):
        stt_paragraph = table.cell(row_idx, 0).paragraphs[0]
        stt_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    note = doc.add_paragraph("Nhóm trưởng: [TODO-TEAM: họ tên] - Ngày: [TODO-TEAM: ngày/tháng/năm]")
    style_paragraph(note, WD_ALIGN_PARAGRAPH.LEFT)
    for run in note.runs:
        style_run(run, size_pt=13, bold=False)


def add_abbreviations_page(doc: Document) -> None:
    add_chapter_title(doc, "DANH MỤC TỪ VIẾT TẮT")

    table = doc.add_table(rows=len(ABBREVIATIONS) + 1, cols=2)
    table.style = "Table Grid"
    table.autofit = False

    for cell in table.columns[0].cells:
        cell.width = Cm(4.0)
    for cell in table.columns[1].cells:
        cell.width = Cm(12.0)

    table.cell(0, 0).text = "Từ viết tắt"
    table.cell(0, 1).text = "Ý nghĩa đầy đủ"
    for col_idx in range(2):
        header_paragraph = table.cell(0, col_idx).paragraphs[0]
        style_paragraph(header_paragraph, WD_ALIGN_PARAGRAPH.CENTER)
        for run in header_paragraph.runs:
            style_run(run, size_pt=13, bold=True)

    for row_idx, (abbr, meaning) in enumerate(ABBREVIATIONS, start=1):
        table.cell(row_idx, 0).text = abbr
        table.cell(row_idx, 1).text = meaning

    style_table_text(table)
    for row_idx in range(1, len(ABBREVIATIONS) + 1):
        table.cell(row_idx, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_figure_list_page(doc: Document) -> None:
    add_chapter_title(doc, "DANH MỤC HÌNH VẼ")
    paragraph = doc.add_paragraph(
        "[TODO-TEAM: Điền số hình và tên sau khi chèn ảnh demo vào Chương 5]"
    )
    style_paragraph(paragraph, WD_ALIGN_PARAGRAPH.LEFT)
    for run in paragraph.runs:
        style_run(run, size_pt=13, bold=False)


def add_toc_placeholder_page(doc: Document) -> None:
    add_chapter_title(doc, "MỤC LỤC")


def add_loi_mo_dau_page(doc: Document, loi_mo_dau_text: str) -> None:
    add_chapter_title(doc, "LỜI MỞ ĐẦU")

    paragraphs = re.split(r"\n\s*\n", loi_mo_dau_text.strip())
    for para_text in paragraphs:
        paragraph = doc.add_paragraph()
        style_paragraph(paragraph, WD_ALIGN_PARAGRAPH.JUSTIFY)

        lines = para_text.splitlines()
        for idx, line in enumerate(lines):
            if idx > 0:
                paragraph.add_run().add_break()
            run = paragraph.add_run(line)
            style_run(run, size_pt=13, bold=False, italic=False)


def build_document(source_md: Path, output_docx: Path) -> None:
    loi_mo_dau_text = extract_loi_mo_dau(source_md)

    doc = Document()
    configure_normal_style(doc)
    configure_section_page(doc.sections[0])

    add_cover_page(doc)
    doc.add_page_break()

    add_task_assignment_page(doc)
    doc.add_page_break()

    add_abbreviations_page(doc)
    doc.add_page_break()

    add_figure_list_page(doc)
    doc.add_page_break()

    add_toc_placeholder_page(doc)

    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section_page(body_section)
    body_section.footer.is_linked_to_previous = False
    set_section_page_number_start(body_section, start=1)

    footer_paragraph = body_section.footer.paragraphs[0]
    footer_paragraph.text = ""
    style_paragraph(footer_paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    add_page_number_field(footer_paragraph)

    add_loi_mo_dau_page(doc, loi_mo_dau_text)

    for section in doc.sections:
        configure_section_page(section)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)


def main() -> None:
    args = parse_args()
    source_md = args.source_md.resolve()
    output_docx = args.output_docx.resolve()
    build_document(source_md=source_md, output_docx=output_docx)
    print("1A DONE - Khung docx và phần đầu hoàn chỉnh.")


if __name__ == "__main__":
    main()
